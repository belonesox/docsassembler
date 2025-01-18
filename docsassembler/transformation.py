# -*- coding: utf-8 -*-
"""
 Translation procedures.
 Name pattern «ext1»2«ext2»
 Where
 * «ext1» — source file extension
 * «ext2» — target file extension
"""
import  os
import sys
import tempfile
import shutil
import subprocess
import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, Mm, RGBColor
# from docx.enum.dml import MSO_THEME_COLOR
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
# import belonesox_tools.MiscUtils as ut

from  lxml import etree
from  .messagefilter import CuteFilter
from .lib import hide_path
from pathlib import Path

CSL_PATH = Path(__file__).parent / 'csl/ieee.csl' 
BIB_MOD_STD = f' --bibliography=bibliography.bib  --citeproc -M link-citations=true --csl {CSL_PATH.as_posix()} '


# pylint: disable-msg=W0612
# :W0612: *Unused variable %r*
#   Used when a variable is defined but not used.
# pylint: disable-msg=W0613
# :W0613: *Unused argument %r*
#   Used when a function or method argument is not used.
# pylint: disable-msg=W0703
# :W0703: *Catch "Exception"*
#   Used when an except catches Exception instances.
# pylint: disable-msg=W0233
# :W0233: *__init__ method from a non direct base class %r is called*
#   Used when an __init__ method is called on a class which is not in the direct

def set_texinput():
    path_to_style_dir = Path(__file__).parent / 'latex'
    old_texpinputs = ''
    if 'TEXINPUTS' in os.environ:
        old_texpinputs = os.environ['TEXINPUTS']
    os.environ['TEXINPUTS'] = ':'.join([old_texpinputs, path_to_style_dir.as_posix()])

def add_page_number(paragraph):
    page_num_run = paragraph.add_run()

    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

def create_footer(doc):
    footer = doc.sections[0].footer
    paragraph = footer.add_paragraph()
    paragraph.alignment = 1  # 1 for center alignment
    add_page_number(paragraph)

def create_header(doc):
    header = doc.sections[0].header
    paragraph = header.add_paragraph()
    paragraph.alignment = 1  # 1 for center alignment
    add_page_number(paragraph)

# def set_cell_border(cell, **kwargs):
#     """
#     Set cell borders using the passed kwargs.
#     """
#     tc = cell._tc
#     tcPr = tc.get_or_add_tcPr()

#     for edge, border_props in kwargs.items():
#         edge_prop = f"tcBorders/{edge}"
#         border = tcPr.find(qn(edge_prop))
#         if border is None:
#             border = OxmlElement(edge_prop)
#             tcPr.append(border)

#         for prop, value in border_props.items():
#             setattr(border, qn(prop), str(value))

def set_cell_border(cell, **kwargs):
    """
    Set cell`s border
    Usage:

    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)

            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def fix_style_for_table(table, doc):
    section = doc.sections[0]
    page_width = section.page_width - section.left_margin - section.right_margin
    table.autofit = False
    table.allow_autofit = False
    table.width = page_width    
    column_width = page_width / (len(table.columns)+1)
    # for i, col in enumerate(table.columns):
    #     if i == 0:
    #         col.width = int(2 * column_width)
    #     else:     
    #         col.width = int(column_width)

    sum_text_row = [0] * len(table.columns)
    for row in table.rows:
        for col_, cell in enumerate(row.cells):
            sum_text_row[col_] += len(cell.text)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": "000000", "space": "0"},
                bottom={"val": "single", "sz": "4", "color": "000000", "space": "0"},
                end={"val": "single", "sz": "4", "color": "000000", "space": "0"},
                start={"val": "single", "sz": "4", "color": "000000", "space": "0"}
            )

    sum_ = sum(sum_text_row)
    column_widths = [column_width/2 + 3*column_width * sum_text_row[col]/sum_ for col in range(len(table.columns))]
    sum_ = sum(column_widths)
    column_widths = [int(page_width/sum_ * col_) for col_ in column_widths]
    for i, col in enumerate(table.columns):
        col.width = column_widths[i]

    # # Set font size for the table cells
    # for row in table.rows:
    #     for cell in row.cells:
    #         for paragraph in cell.paragraphs:
    #             for run in paragraph.runs:
    #                 run.font.size = Pt(12)

def tune_docx(document):
    def sv():
        document.save('wtf.docx')
    # create_footer(document)        
    create_header(document)        
    sv()
    section = document.sections[0]
    section.page_height = Mm(297)  
    section.page_width = Mm(210)   
    section.left_margin = Mm(25.4/2) 
    section.right_margin = Mm(25.4/2)
    section.top_margin = Mm(25.4/2)  
    section.bottom_margin = Mm(25.4/2)
    section.header_distance = Mm(12.7) 
    section.footer_distance = Mm(12.7) 
    for style_ in document.styles:
        print(style_.name)
        if 'font' in dir(style_):
            if style_.name in [ #'Normal', 
                                'Text Body', 'Body Text']:
                style_.font.size = Pt(12)
                style_.font.name = 'Cambria'
                style_.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif style_.name in ['Source Code']:
                style_.font.size = Pt(8)
                style_.font.name = 'Consolas'
            elif 'Heading' in style_.name:
                style_.font.name = 'Calibri'
                sz = 0 
                try:
                    sz = int(style_.name[-1:])
                except:
                    ...    
                style_.font.size = Pt(12 + int( (6-sz) ))
                style_.font.color.rgb = RGBColor(0, 0, 0)
                ...
        ...
        for i, table_ in enumerate(document.tables):
            fix_style_for_table(table_, document)
            # print(i, dir(table_))

    sv()
    ...    

class Transformation:

    @staticmethod
    def dot2pdf_generic(env, target, source):
        """
            .DOT —> .PDF
        """  
        dot = "dot"
        if os.path.splitext(source[0].abspath)[1] == ".neato":
            dot = "neato"
        if os.path.splitext(source[0].abspath)[1] == ".fdp":
            dot = "fdp"
        if os.path.splitext(source[0].abspath)[1] == ".circo":
            dot = "circo"
        src = source[0].abspath
        tmp = target[0].abspath + ".svg"
        (pathname, ext) = os.path.splitext(target[0].abspath)

        dot = os.path.join(env.project_db['paths']['graphviz'], dot)
        os.system('%(dot)s -Tsvg %(src)s -Gbgcolor=transparent -Gmargin=0 -o %(tmp)s' % vars())
        tgt = target[0].abspath
        inkscapepath = os.path.join(env.project_db['paths']['inkscape'], 'inkscape')
        command = ''.join([inkscapepath, ' --without-gui --export-area-drawing ',
                        ' --export-background-opacity=0 --export-pdf="%s.pdf" ',
                        ' --export-text-to-path "%s"']) % (pathname, tmp)
        os.system(command)


    @staticmethod
    def dot2pdf(env, target, source):
        """
        Translate DOT files to PDF
        """
        Transformation.dot2pdf_generic(env, target, source)

    @staticmethod
    def neato2pdf(env, target, source):
        """
        Translate NEATO files to PDF
        """
        Transformation.dot2pdf_generic(env, target, source)

    @staticmethod
    def fdp2pdf(env, target, source):
        """
        Translate FDP files to PDF
        """
        Transformation.dot2pdf_generic(env, target, source)

    @staticmethod
    def circo2pdf(env, target, source):
        """
        Translate CIRCO files to PDF
        """
        Transformation.dot2pdf_generic(env, target, source)

    @staticmethod
    def dot2svg_generic(env, target, source):
        """
        Translate DOT files to SVG (generic procedure)
        """
        dot = "dot"
        ext = os.path.splitext(source[0].abspath)[1]
        if ext in [".neato", ".fdp", ".circo"]:
            dot = ext[1:] 
        src = source[0].abspath
        tgt = target[0].abspath

        dot = os.path.join(env.project_db['paths']['graphviz'], dot)
        os.system('%(dot)s -Tsvg %(src)s -Gbgcolor=transparent -Gmargin=0 -o %(tgt)s' % vars())

    @staticmethod
    def dot2svg(env, target, source):
        """
        Translate DOT files to SVG
        """
        Transformation.dot2svg_generic(env, target, source)

    @staticmethod
    def neato2svg(env, target, source):
        """
        Translate NEATO files to SVG
        """
        Transformation.dot2svg_generic(env, target, source)

    @staticmethod
    def fdp2svg(env, target, source):
        """
        Translate FDP files to SVG
        """
        Transformation.dot2svg_generic(env, target, source)

    @staticmethod
    def circo2svg(env, target, source):
        """
        Translate CIRCO files to SVG
        """
        Transformation.dot2svg_generic(env, target, source)

    @staticmethod
    def twopi2svg(env, target, source):
        """
        Translate TWOPI files to SVG
        """
        Transformation.dot2svg_generic(env, target, source)

    @staticmethod
    def svg2eps(target, source, env):
        """
        Translate SVG files to EPS
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        assert(ext in [".eps"])
        inkscapepath = os.path.join(env.project_db['paths']['inkscape'], 'inkscape')
        command = ''.join([inkscapepath, ' --without-gui --export-area-drawing ',
                    ' --export-background-opacity=0 --export-eps="%s.eps" ',
                    ' --export-text-to-path "%s"']) % (pathname, source[0].abspath)
        os.system(command)

    @staticmethod
    def md2tex(target, source, env):
        """
        Translate Markdown files to Tex
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        template_path = Path(__file__).parent / 'latex/docstruct.latex' 
        assert(ext in [".tex"])

        bib_mod = ''
        if Path('bibliography.bib').exists():
            bib_mod = BIB_MOD_STD

        scmd = f'''
pandoc -s   {source[0].abspath} {bib_mod} --template {template_path.as_posix()}  -o   {target[0].abspath} 
'''.replace('\n', ' ').strip()

        os.system(scmd)


    @staticmethod
    def md2html(target, source, env):
        """
        Translate Markdown files to HTML
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        os.chdir(Path(target[0].abspath).parent)
        metadata_path = Path(__file__).parent / 'markdown/metadata.yaml' 
        pd_filter_path = Path(__file__).parent / 'pandoc/filters' 
        pd_readers_path = Path(__file__).parent / 'pandoc/readers' 

        s5_path = Path(__file__).parent / 's5' 
        math_path = Path(__file__).parent / 'math/tex-chtml-full.js' 
        # from_mod = ' --from gfm'
        pandoc_source_format = 'gfm+definition_lists+sourcepos'
        os.environ['PANDOC_SOURCE_FORMAT'] = pandoc_source_format
        from_mod = f' --from {pandoc_source_format}'
        # from_mod = ' --from commonmark_x+sourcepos'
        # from_mod = ' --from commonmark_x'
        # from_mod = ' --from gfm+definition_lists '
        # from_mod = f' --variable dasws_preview  --from {pd_readers_path}/commonmark_x.lua+sourcepos+definition_lists '
        
        title_exists = False
        with open(source[0].abspath, 'r', encoding='utf-8') as lf:
            input_text_ = lf.read()
            title_exists = 'title:' in input_text_

        pd_filter_path = Path(__file__).parent / 'pandoc/filters' 

        dasws_mod_ = ''
        # dasws_mod_ = f' --lua-filter="{pd_filter_path}/sourcepos_sync.lua" --lua-filter="{pd_filter_path}/dasws_output.lua"  '

        title_mod_ = ''
        if not title_exists:
            terms_ = source[0].abspath.split(os.path.sep)[-4:-1]
            terms_.reverse()
            title_ = ' / '.join(terms_)     
            title_mod_ = f'--metadata title="{title_}" '

        template_path = Path(__file__).parent / 'pandoc/templates/html/default.html5'

        slides_mod = ''
        if '.slides.' in source[0].abspath:
            from_mod = ''
            # s5_template_path = f'{s5_path}/our-s5.html' 
            template_path = f'{s5_path}/our-s5.html' 
            s5_path = './s5/'
            slides_mod = f' --to s5  -V s5-url={s5_path}/our '

        bib_mod = ''
        if Path('bibliography.bib').exists():
            bib_mod = BIB_MOD_STD
            from_mod = f' --from markdown+citations+definition_lists '

        template_mod = f'--template {template_path}'    
        assert(ext in [".html"])
        text = Path(source[0].abspath).read_text()
        (Path(target[0].abspath).parent / '.texpics').mkdir(exist_ok=True, parents=True)
        toc_mod = ''
        if '<!--- TOC --->' in text:
            toc_mod = '--toc'
# pandoc -s input_file.md -t json | gladtex -P -  | pandoc -s -f json
        # pandoc -t json -s {from_mod} --lua-filter="{pd_filter_path}/include-files.lua --filter pandoc-include   "
        scmd = f'''
        pandoc -t json -s {from_mod} --filter pandoc-include   
        --metadata-file "{metadata_path}" {bib_mod} "{source[0].abspath}" 
        | gladtex -P -K -d .texpics -c 0019F7 - | 
        pandoc {slides_mod} {title_mod_} {dasws_mod_} {template_mod} {bib_mod} --output "{target[0].abspath}"  
        --standalone --embed-resources {toc_mod} -f json
        '''.replace('\n', ' ').strip()
        #--embed-resources
        Path('debug-all.sh').write_text(scmd)
        os.system(scmd)

        scmd = f'''
        pandoc -t json -s {from_mod} --filter pandoc-include   
        --metadata-file "{metadata_path}"  "{source[0].abspath}" 
        | jq . > debug.json
        '''.replace('\n', ' ').strip()
        Path('debug-1st-phase.sh').write_text("#!/bin/sh\n" + scmd)

        scmd = f'''
        cat debug.json 
        | gladtex -P -K -d .texpics -c 0019F7 - | 
        pandoc {slides_mod} {title_mod_} {dasws_mod_} --output "{target[0].abspath}"  
        --standalone --embed-resources {toc_mod} -f json
        '''.replace('\n', ' ').strip()
        Path('debug-2nd-phase.sh').write_text("#!/bin/sh\n" + scmd)


        # scmd = f'''pandoc {slides_mod} --embed-resources --gladtex  --filter pandoc-include  --standalone --output ".tmp.htex" --metadata-file "{metadata_path}"  "{source[0].abspath}" '''
        # # print(scmd)
        # os.system(scmd)
        # scmd = f'''gladtex -o {target[0].abspath} -d texpics .tmp.htex'''
        # # scmd = f'''pandoc {slides_mod} --embed-resources --gladtex  --filter pandoc-include  --standalone --output "{target[0].abspath}" --metadata-file "{metadata_path}"  "{source[0].abspath}" '''
        # #--embed-resources
        # #--mathjax={math_path}
        # os.system(scmd)


    @staticmethod
    def md2docx(target, source, env):
        """
        Translate Markdown files to HTML
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        # metadata_path = Path(__file__).parent / 'markdown/metadata.yaml' 
        s5_path = Path(__file__).parent / 's5' 
        math_path = Path(__file__).parent / 'math/tex-chtml-full.js' 
        from_mod = ' --from gfm+definition_lists  '
        to_mod = ' --to html '
        title_exists = False
        with open(source[0].abspath, 'r', encoding='utf-8') as lf:
            input_text_ = lf.read()
            title_exists = 'title:' in input_text_
 
        bib_mod = ''
        if Path('bibliography.bib').exists():
            bib_mod = BIB_MOD_STD
            from_mod = f' --from markdown+citations+definition_lists '

        title_mod_ = ''
        if not title_exists:
            terms_ = source[0].abspath.split(os.path.sep)[-4:-1]
            terms_.reverse()
            title_ = ' / '.join(terms_)     
            title_mod_ = f'--metadata title="{title_}" '


        reference_mod = ''
        reference_doc = Path(target[0].abspath).parent / '.das/templates/docx/reference.docx'
        if reference_doc.exists():
            reference_mod = f' --reference-doc={reference_doc} '

        assert(ext in [".docx"])
        text = Path(source[0].abspath).read_text()
        (Path(target[0].abspath).parent / '.texpics').mkdir(exist_ok=True, parents=True)
        toc_mod = ''
        if '<!--- TOC --->' in text:
            toc_mod = '--toc -M toc-title="Содержание" '
# pandoc -s input_file.md -t json | gladtex -P -  | pandoc -s -V lang:ru -f json
        scmd = f'''
        pandoc -t json -s {from_mod} {bib_mod} --filter pandoc-include   
            "{source[0].abspath}" 
        |
        pandoc {title_mod_} --output "{target[0].abspath}"  
        --standalone --embed-resources {toc_mod} {reference_mod} -V lang:ru -f json
        '''.replace('\n', ' ').strip()
        print(scmd)
        Path('debug-docx.sh').write_text("#!/bin/sh\n" + scmd)
        os.system(scmd)
        # | gladtex -P -K -d .texpics -c 0019F7 - | 

        document = docx.Document(target[0].abspath)
        tune_docx(document)
        document.save(target[0].abspath)



    @staticmethod
    def das2docx(target, source, env):
        """
        Translate DAS files to HTML
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        document = docx.Document(target[0].abspath)
        tune_docx(document)

        import yaml
        das_ = yaml.unsafe_load(open(source[0].abspath, 'r'))
        from_mod = ' --from json  '
        to_mod = ' --to docx '

        title_mod_ = ''
        if 'title' in das_:
            title_mod_ = f'''--metadata title="{das_['title']}" '''

        input_files = []
        curdir = os.getcwd()
        for input_ in das_['files']:
            os.chdir(curdir)
            basedir, infilename = os.path.split(input_)
            os.chdir(basedir)
            scmd = f'das {infilename}'
            os.system(scmd)
            input_files.append(input_)
        input_files_mod = " ".join(input_files)    

        os.chdir(curdir)
        reference_mod = ''
        reference_doc = Path(target[0].abspath).parent / '.das/templates/docx/reference.docx'
        if reference_doc.exists():
            reference_mod = f' --reference-doc={reference_doc} '

        assert(ext in [".docx"])
        # text = Path(source[0].abspath).read_text()
        (Path(target[0].abspath).parent / '.texpics').mkdir(exist_ok=True, parents=True)
        toc_mod = '--toc -M toc-title="Содержание" '
        scmd = f'''
        pandoc -t json -s {from_mod} --filter pandoc-include   
            {input_files_mod}
        |
        pandoc {title_mod_} --output "{target[0].abspath}"  
        --standalone --embed-resources {toc_mod} {reference_mod} -V lang:ru -f json
        '''.replace('\n', ' ').strip()
        print(scmd)
        Path('debug-docx.sh').write_text("#!/bin/sh\n" + scmd)
        os.system(scmd)

        document = docx.Document(target[0].abspath)
        tune_docx(document)
        document.save(target[0].abspath)

        # | gladtex -P -K -d .texpics -c 0019F7 - | 



    @staticmethod
    def md2pandoc(target, source, env):
        """
        Translate Markdown files to Internal Pandoc format
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        os.chdir(Path(target[0].abspath).parent)
        metadata_path = Path(__file__).parent / 'markdown/metadata.yaml' 
        pd_filter_path = Path(__file__).parent / 'pandoc/filters' 
        pd_readers_path = Path(__file__).parent / 'pandoc/readers' 

        s5_path = Path(__file__).parent / 's5' 
        math_path = Path(__file__).parent / 'math/tex-chtml-full.js' 
        # from_mod = ' --from gfm'
        pandoc_source_format = 'gfm+definition_lists+sourcepos'
        os.environ['PANDOC_SOURCE_FORMAT'] = pandoc_source_format
        from_mod = f' --from {pandoc_source_format}'
        
        title_exists = False
        with open(source[0].abspath, 'r', encoding='utf-8') as lf:
            input_text_ = lf.read()
            title_exists = 'title:' in input_text_

        pd_filter_path = Path(__file__).parent / 'pandoc/filters' 

        dasws_mod_ = ''
        # dasws_mod_ = f' --lua-filter="{pd_filter_path}/sourcepos_sync.lua" --lua-filter="{pd_filter_path}/dasws_output.lua"  '

        title_mod_ = ''
        if not title_exists:
            terms_ = source[0].abspath.split(os.path.sep)[-4:-1]
            terms_.reverse()
            title_ = ' / '.join(terms_)     
            title_mod_ = f'--metadata title="{title_}" '

        assert(ext in [".pandoc"])
        text = Path(source[0].abspath).read_text()
        (Path(target[0].abspath).parent / '.texpics').mkdir(exist_ok=True, parents=True)
        toc_mod = ''
        if '<!--- TOC --->' in text:
            toc_mod = '--toc'
        scmd = f'''
        pandoc -t json -s {from_mod} --filter pandoc-include   
        --metadata-file "{metadata_path}"  "{source[0].abspath}" 
        {title_mod_} {dasws_mod_} --output "{target[0].abspath}"  
        --standalone --embed-resources {toc_mod} 
        '''.replace('\n', ' ').strip()
        # Path('debug-pandoc.sh').write_text("#!/bin/sh\n" + scmd)
        os.system(scmd)


    @staticmethod
    def html2pandoc(target, source, env):
        """
        Translate HTML files to Internal Pandoc format
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        os.chdir(Path(target[0].abspath).parent)
        metadata_path = Path(__file__).parent / 'markdown/metadata.yaml' 
        from_mod = f' --from html'
        
        title_exists = False
        with open(source[0].abspath, 'r', encoding='utf-8') as lf:
            input_text_ = lf.read()
            title_exists = 'title:' in input_text_

        title_mod_ = ''
        if not title_exists:
            terms_ = source[0].abspath.split(os.path.sep)[-4:-1]
            terms_.reverse()
            title_ = ' / '.join(terms_)     
            title_mod_ = f'--metadata title="{title_}" '

        assert(ext in [".pandoc"])
        text = Path(source[0].abspath).read_text()
        (Path(target[0].abspath).parent / '.texpics').mkdir(exist_ok=True, parents=True)
        toc_mod = ''
        if '<!--- TOC --->' in text:
            toc_mod = '--toc'
        scmd = f'''
        pandoc -t json -s {from_mod} --filter pandoc-include   
        --metadata-file "{metadata_path}"  "{source[0].abspath}" 
        {title_mod_} --output "{target[0].abspath}"  
        --standalone --embed-resources {toc_mod} 
        '''.replace('\n', ' ').strip()
        # Path('debug-pandoc.sh').write_text("#!/bin/sh\n" + scmd)
        os.system(scmd)


    @staticmethod
    def html2docx(target, source, env):
        """
        Translate HTML files to DOCX
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        from_mod = ' --from html  '
        to_mod = ' --to docx '
        title_exists = False
        with open(source[0].abspath, 'r', encoding='utf-8') as lf:
            input_text_ = lf.read()
            title_exists = 'title:' in input_text_
 
        title_mod_ = ''
        if not title_exists:
            terms_ = source[0].abspath.split(os.path.sep)[-4:-1]
            terms_.reverse()
            title_ = ' / '.join(terms_)     
            title_mod_ = f'--metadata title="{title_}" '


        reference_mod = ''
        reference_doc = Path(target[0].abspath).parent / '.das/templates/docx/reference.docx'
        if reference_doc.exists():
            reference_mod = f' --reference-doc={reference_doc} '

        assert(ext in [".docx"])
        text = Path(source[0].abspath).read_text()
        (Path(target[0].abspath).parent / '.texpics').mkdir(exist_ok=True, parents=True)
        toc_mod = ''
        if '<!--- TOC --->' in text:
            toc_mod = '--toc -M toc-title="Содержание" '
# pandoc -s input_file.md -t json | gladtex -P -  | pandoc -s -V lang:ru -f json
        scmd = f'''
        pandoc -t json -s {from_mod} --filter pandoc-include   
            "{source[0].abspath}" 
        |
        pandoc {title_mod_} --output "{target[0].abspath}"  
        --standalone --embed-resources {toc_mod} {reference_mod} -V lang:ru -f json
        '''.replace('\n', ' ').strip()
        print(scmd)
        Path('debug-docx.sh').write_text("#!/bin/sh\n" + scmd)
        os.system(scmd)


    @staticmethod
    def tex2pdf(target, source, env):
        """
        Translate SVG files to EPS
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        assert(ext in [".pdf"])
        tmpname = os.path.join(tempfile.gettempdir(), tempfile.gettempprefix())
        tmptexname = tmpname + ".tex"
        tmppdfname = tmpname + ".pdf"
        tmppcropname = tmpname + ".crop.pdf"
        shutil.copy(source[0].abspath, tmptexname)
        set_texinput()
        scmd = os.path.join(env.project_db['paths']['tex'],
            'xelatex -interaction nonstopmode "%(tmptexname)s"'
            % vars())
        curdir = os.getcwd()
        os.chdir(tempfile.gettempdir())
        
        cutefilter = CuteFilter(tempfile.gettempdir())
        out= "Mock:" + scmd
        try:
            out = subprocess.check_output(scmd, shell=True).decode("utf8")    
        except Exception as ex_:    
            print(ex_)

        out = cutefilter(out)
        try:
            os.system(scmd)
        except Exception as ex_:    
            print(ex_)

        print(os.environ['PATH'])
        print("!!", scmd)
        scmd = os.path.join(env.project_db['paths']['tex'],
            "".join(['pdfcrop "', tmppdfname, '" "', target[0].abspath, '"']))
        os.system(scmd)
        os.chdir(curdir)
        print("!tex2pdf OK!")

    @staticmethod
    def preprocess_svg(sourcefile):
        """
        Preprocess possible composite SVG file,
        for all nodes with label like "xxx.svg"
        include "xxx.svg" instead of the node.
        """
        lf = open(sourcefile, 'r', encoding='utf-8')   
        doc = etree.parse(lf)
        lf.close()
        root = doc.getroot() 
        for element in root.iterfind(".//{http://www.w3.org/2000/svg}g"):
            if "{http://www.inkscape.org/namespaces/inkscape}label" in element.attrib:
                label = element.get("{http://www.inkscape.org/namespaces/inkscape}label").strip()
                if label.endswith(".svg"):
                    includedsvg = os.path.realpath(os.path.join(os.path.split(sourcefile)[0], label))
                    if os.path.exists(includedsvg):
                        ls = Transformation.preprocess_svg(includedsvg)
                        svgroot = etree.fromstring(ls)
                        for child in element:
                            element.remove(child)
                        #if 'width' in svgroot.attrib:
                        #   del svgroot.attrib['width']
                        #if 'height' in svgroot.attrib:
                        #   del svgroot.attrib['height']
                        svgroot.attrib['width'] = r"100%"  
                        svgroot.attrib['height'] = r"100%"  
    #                 if 'viewBox' in svgroot.attrib:
    #                    del svgroot.attrib['viewBox']
                        element.append(svgroot)
        res_ =  etree.tostring(root, pretty_print=True, encoding = "unicode")    
        # res_ = res_.decode('utf-8')
        return res_


    @staticmethod
    def svg2pdf(target, source, env):
        """
        Translate SVG files to PDF
        """
        (pathname, ext) = os.path.splitext(target[0].abspath)
        assert(ext in [".pdf"])
        preprocessedname = source[0].abspath
        ls = Transformation.preprocess_svg(source[0].abspath)
        preprocessedname = os.path.join(os.path.split(source[0].abspath)[0],
                                        os.path.split(source[0].abspath)[1] + ".obj")
        if os.path.exists(preprocessedname):
            os.unlink(preprocessedname)
        lf = open(preprocessedname, "w", encoding='utf-8')   
        lf.write(ls)
        lf.close()
        hide_path(preprocessedname)
        inkpath=env.project_db['paths']['inkscape']
        command = os.path.join(inkpath,
                    ''.join(['inkscape --without-gui --export-area-drawing ',
                        ' --export-background-opacity=0 --export-pdf="%s.pdf" ',
                        ' --export-text-to-path "%s"']) % (pathname, preprocessedname) )
        os.system(command)

